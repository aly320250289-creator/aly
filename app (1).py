#!/usr/bin/env python3
"""
================================================================================
ARCHITECTURE AI PLATFORM - STANDALONE DESKTOP APPLICATION
================================================================================
Single-file solution: Backend + Frontend + Auto-installer + Browser launcher
Extracts structural/architectural data from PDF/Word using OpenAI GPT-4o-mini

Usage:
    python app.py              # Development mode
    # OR double-click .exe     # Production mode (after PyInstaller)

Author: Architecture AI Platform
Version: 1.0.0
================================================================================
"""

import os
import sys
import json
import time
import threading
import webbrowser
from pathlib import Path

# =============================================================================
# SECTION 1: AUTO-DEPENDENCY INSTALLER
# =============================================================================
REQUIRED_PACKAGES = {
    'flask': 'Flask>=2.3.0',
    'openai': 'openai>=1.0.0',
    'pdfplumber': 'pdfplumber>=0.10.0',
    'python_docx': 'python-docx>=0.8.11',
    'werkzeug': 'Werkzeug>=2.3.0',
}

def check_and_install_dependencies():
    """Check if required packages are installed, auto-install if missing."""
    missing = []
    for module, package in REQUIRED_PACKAGES.items():
        try:
            __import__(module.replace('_', '-'))
        except ImportError:
            missing.append(package)

    if missing:
        print("\n" + "="*60)
        print("📦 MISSING DEPENDENCIES DETECTED")
        print("="*60)
        print(f"   Required: {', '.join(missing)}")
        print("\n⏳ Installing now (this may take 1-2 minutes)...")
        print("="*60 + "\n")

        import subprocess
        try:
            subprocess.check_call(
                [sys.executable, "-m", "pip", "install", "--quiet", *missing],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL
            )
            print("✅ All dependencies installed successfully!\n")
            print("🔄 Restarting application to load new packages...\n")
            time.sleep(1)
            os.execv(sys.executable, [sys.executable] + sys.argv)
        except subprocess.CalledProcessError as e:
            print(f"\n❌ FAILED to install dependencies: {e}")
            print("\nPlease run manually:")
            print(f"   pip install {' '.join(missing)}")
            sys.exit(1)

check_and_install_dependencies()

# =============================================================================
# SECTION 2: IMPORTS (After dependency check)
# =============================================================================
from flask import Flask, render_template_string, request, jsonify
from werkzeug.utils import secure_filename
import pdfplumber
from docx import Document
import openai

# =============================================================================
# SECTION 3: CONFIGURATION & CONSTANTS
# =============================================================================
APP_NAME = "Architecture AI Platform"
APP_VERSION = "1.0.0"
APP_PORT = 5000
APP_HOST = "0.0.0.0"

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max upload

# Paths
BASE_DIR = Path(__file__).parent.resolve()
UPLOAD_FOLDER = BASE_DIR / "temp_uploads"
UPLOAD_FOLDER.mkdir(exist_ok=True)

# Supported file types
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc'}

# OpenAI API Key (from environment variable)
OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY', '')

# =============================================================================
# SECTION 4: FILE READING FUNCTIONS
# =============================================================================
def allowed_file(filename):
    """Check if file extension is supported."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    """Extract text from PDF using pdfplumber."""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- PAGE {i} ---\n{page_text}\n"
    except Exception as e:
        return f"ERROR_PDF: {str(e)}"
    return text.strip()

def extract_text_from_docx(file_path):
    """Extract text from Word document using python-docx."""
    text = ""
    try:
        doc = Document(file_path)
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        # Extract tables
        for table in doc.tables:
            text += "\n--- TABLE ---\n"
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    text += row_text + "\n"
    except Exception as e:
        return f"ERROR_DOCX: {str(e)}"
    return text.strip()

def extract_text(file_path, filename):
    """Route to appropriate extractor based on file extension."""
    ext = filename.rsplit('.', 1)[1].lower()
    if ext == 'pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['docx', 'doc']:
        return extract_text_from_docx(file_path)
    else:
        return "ERROR: Unsupported file type"

# =============================================================================
# SECTION 5: OPENAI HYBRID EXTRACTION LOGIC
# =============================================================================
def build_system_prompt(user_focus=""):
    """
    HYBRID EXTRACTION LOGIC:

    SCENARIO A (DEFAULT): If user_focus is empty → Extract Architectural Basics
    SCENARIO B (CUSTOM): If user_focus has text → Deep extract ONLY that focus
    """

    if not user_focus or user_focus.strip() == "":
        # ═══════════════════════════════════════════════════════════════
        # SCENARIO A: DEFAULT ARCHITECTURAL BASICS EXTRACTION
        # ═══════════════════════════════════════════════════════════════
        return """You are an expert Architectural and Construction Document Analyzer.
Your task is to extract "Architectural Basics" from construction documents and return ONLY valid JSON.

Extract the following structured information:

{
  "project_name": "string or null - The official project name from title block",
  "client": "string or null - Client/Owner name",
  "contractor": "string or null - General contractor or builder",
  "architect": "string or null - Architecture firm or architect name",
  "engineer": "string or null - Structural engineer or engineering firm",
  "project_dates": {
    "start_date": "string or null (MM/DD/YYYY format)",
    "completion_date": "string or null (MM/DD/YYYY format)",
    "drawing_date": "string or null (MM/DD/YYYY format) - Date on drawing sheet"
  },
  "main_materials": [
    "List of primary construction materials: concrete grade, steel type, lumber species, ",
    "insulation type, roofing material, flooring, windows, doors, etc."
  ],
  "budget_info": {
    "total_budget": "string or null - Total project cost if stated",
    "currency": "string or null - USD, EUR, etc.",
    "cost_breakdown": "string or null - Any cost breakdown or unit prices mentioned"
  },
  "structural_info": {
    "building_type": "string or null - Residential, commercial, industrial, etc.",
    "foundation_type": "string or null - Slab, crawl space, basement, pier, etc.",
    "framing_type": "string or null - Wood frame, steel frame, concrete, masonry, etc.",
    "roof_type": "string or null - Gable, hip, flat, truss, etc."
  },
  "location": {
    "address": "string or null - Street address",
    "city": "string or null",
    "state": "string or null",
    "zip_code": "string or null"
  },
  "drawing_info": {
    "sheet_number": "string or null - Drawing sheet number (e.g., S1.1, A-101)",
    "drawing_scale": "string or null - Scale shown on drawing (e.g., 1/4" = 1'-0")",
    "revision": "string or null - Revision number or date"
  }
}

RULES:
- Return ONLY valid JSON. No markdown, no explanations, no code blocks.
- Use null for fields not found in the document. Do NOT use empty strings.
- For dates, prefer MM/DD/YYYY format. If only partial date found, include as-is.
- For materials, extract specific grades and types (e.g., "3000 PSI concrete", "Southern Pine #2").
- For budget, look for total costs, unit prices, allowances, or cost references.
- For structural_info, infer from context if not explicitly stated (e.g., "wood frame" from lumber mentions).
- drawing_info comes from the title block or sheet header.
- Be precise - use exact text from the document when possible."""

    else:
        # ═══════════════════════════════════════════════════════════════
        # SCENARIO B: CUSTOM FOCUS DEEP EXTRACTION
        # ═══════════════════════════════════════════════════════════════
        return f"""You are an expert Architectural and Construction Document Analyzer.
The user has requested a SPECIFIC FOCUS for extraction. You must IGNORE all other information.

USER'S SPECIFIC REQUEST: "{user_focus}"

Your task is to DEEPLY extract ONLY information related to the user's request.
Search through ALL pages, tables, notes, schedules, and specifications.

Return valid JSON with this exact structure:
{{
  "extraction_focus": "Brief description of what was requested",
  "found_items": [
    {{
      "item": "Specific detail or item found",
      "value": "The exact value, specification, or measurement",
      "context": "Brief surrounding text for context (2-3 sentences)",
      "location_hint": "Where in document found (e.g., 'Page 2, Structural Notes', 'Table: Concrete Schedule')"
    }}
  ],
  "summary": "Concise summary of all findings related to the user's focus",
  "confidence": "high OR medium OR low - based on clarity and completeness of information",
  "additional_notes": "Any important caveats, related specs, or missing information the user should know"
}}

RULES:
- Return ONLY valid JSON. No markdown, no explanations.
- Extract EVERY instance related to "{user_focus}". Do not stop at the first match.
- Include exact specifications: grades, dimensions, quantities, spacing, depths, strengths.
- If the requested info is NOT found, return empty found_items array and explain in summary.
- For numerical values, include units EXACTLY as stated (e.g., "6\" OC", "3000 PSI", "2x10").
- Check tables, schedules, notes, and specifications thoroughly."""

def extract_with_openai(document_text, user_focus=""):
    """Send document text to OpenAI and return structured JSON extraction."""

    if not OPENAI_API_KEY:
        return {
            "error": "OpenAI API Key not configured. Set the OPENAI_API_KEY environment variable.",
            "setup_help": "Windows: set OPENAI_API_KEY=sk-your-key\nMac/Linux: export OPENAI_API_KEY=sk-your-key"
        }

    # Truncate if too long (stay within token limits, keep some room for response)
    max_chars = 120000
    if len(document_text) > max_chars:
        document_text = document_text[:max_chars] + "\n\n[Document truncated due to length...]"

    system_prompt = build_system_prompt(user_focus)

    try:
        client = openai.OpenAI(api_key=OPENAI_API_KEY)

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system", 
                    "content": system_prompt
                },
                {
                    "role": "user", 
                    "content": f"Analyze this construction document and extract the requested information:\n\n{document_text}"
                }
            ],
            temperature=0.1,  # Low temperature for consistent, factual extraction
            max_tokens=4000,
            response_format={"type": "json_object"}
        )

        # Parse JSON response
        result_text = response.choices[0].message.content
        result_json = json.loads(result_text)

        # Add metadata
        result_json["_meta"] = {
            "app_name": APP_NAME,
            "app_version": APP_VERSION,
            "model_used": "gpt-4o-mini",
            "extraction_mode": "custom_focus" if user_focus.strip() else "default_basics",
            "user_focus": user_focus if user_focus.strip() else None,
            "document_length_chars": len(document_text),
            "tokens_used": response.usage.total_tokens if response.usage else None,
            "prompt_tokens": response.usage.prompt_tokens if response.usage else None,
            "completion_tokens": response.usage.completion_tokens if response.usage else None
        }

        return result_json

    except openai.AuthenticationError:
        return {"error": "Invalid OpenAI API Key. Please check your key at https://platform.openai.com/api-keys"}
    except openai.RateLimitError:
        return {"error": "OpenAI rate limit exceeded. Please wait a moment and try again."}
    except openai.APIConnectionError:
        return {"error": "Cannot connect to OpenAI API. Check your internet connection."}
    except json.JSONDecodeError as e:
        return {"error": f"AI returned invalid JSON. Please try again. Details: {str(e)}"}
    except Exception as e:
        return {"error": f"OpenAI API error: {str(e)}"}

# =============================================================================
# SECTION 6: HTML FRONTEND (Embedded as Python string)
# =============================================================================
HTML_TEMPLATE = r"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Architecture AI Platform - Document Analyzer</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <link href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css" rel="stylesheet">
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap" rel="stylesheet">
    <style>
        body { font-family: 'Inter', sans-serif; }
        .drop-zone {
            border: 3px dashed #cbd5e1;
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        }
        .drop-zone:hover {
            border-color: #94a3b8;
            background-color: #f8fafc;
        }
        .drop-zone.dragover {
            border-color: #3b82f6;
            background-color: #eff6ff;
            transform: scale(1.02);
            box-shadow: 0 0 0 4px rgba(59, 130, 246, 0.1);
        }
        .drop-zone.has-file {
            border-color: #10b981;
            background-color: #ecfdf5;
        }
        .spinner {
            border: 4px solid #f1f5f9;
            border-top: 4px solid #3b82f6;
            border-radius: 50%;
            width: 48px;
            height: 48px;
            animation: spin 1s linear infinite;
        }
        @keyframes spin {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }
        .result-card {
            animation: slideUp 0.5s cubic-bezier(0.16, 1, 0.3, 1);
        }
        @keyframes slideUp {
            from { opacity: 0; transform: translateY(24px); }
            to { opacity: 1; transform: translateY(0); }
        }
        .pulse-ring {
            animation: pulse-ring 2s cubic-bezier(0.215, 0.61, 0.355, 1) infinite;
        }
        @keyframes pulse-ring {
            0% { transform: scale(0.8); opacity: 0.8; }
            100% { transform: scale(2); opacity: 0; }
        }
    </style>
</head>
<body class="bg-slate-50 min-h-screen">
    <!-- Header -->
    <header class="bg-slate-900 text-white shadow-2xl">
        <div class="max-w-6xl mx-auto px-6 py-5">
            <div class="flex items-center justify-between">
                <div class="flex items-center gap-4">
                    <div class="relative">
                        <div class="w-11 h-11 bg-blue-600 rounded-xl flex items-center justify-center shadow-lg shadow-blue-500/30">
                            <i class="fas fa-building text-lg"></i>
                        </div>
                        <div class="absolute inset-0 rounded-xl bg-blue-400 pulse-ring -z-10"></div>
                    </div>
                    <div>
                        <h1 class="text-xl font-bold tracking-tight">Architecture AI Platform</h1>
                        <p class="text-slate-400 text-xs font-medium">Construction Document Intelligence &bull; GPT-4o-mini</p>
                    </div>
                </div>
                <div class="flex items-center gap-3">
                    <span id="apiStatus" class="px-3 py-1.5 rounded-full text-xs font-semibold bg-amber-500/15 text-amber-400 border border-amber-500/20 flex items-center gap-1.5">
                        <span class="w-1.5 h-1.5 rounded-full bg-amber-400 animate-pulse"></span>
                        Checking API...
                    </span>
                </div>
            </div>
        </div>
    </header>

    <main class="max-w-6xl mx-auto px-6 py-8">
        <!-- Upload Section -->
        <div class="bg-white rounded-2xl shadow-xl shadow-slate-200/50 p-8 mb-6 border border-slate-100">
            <div class="flex items-center gap-3 mb-6">
                <div class="w-10 h-10 bg-blue-50 rounded-xl flex items-center justify-center">
                    <i class="fas fa-cloud-upload-alt text-blue-600"></i>
                </div>
                <div>
                    <h2 class="text-lg font-bold text-slate-800">Upload Document</h2>
                    <p class="text-sm text-slate-500">Drop your PDF or Word construction document</p>
                </div>
            </div>

            <!-- Drag & Drop Zone -->
            <div id="dropZone" class="drop-zone rounded-2xl p-12 text-center cursor-pointer mb-6 relative overflow-hidden">
                <input type="file" id="fileInput" accept=".pdf,.docx,.doc" class="hidden">

                <!-- Background decoration -->
                <div class="absolute inset-0 opacity-[0.03] pointer-events-none">
                    <svg width="100%" height="100%">
                        <pattern id="grid" width="40" height="40" patternUnits="userSpaceOnUse">
                            <path d="M 40 0 L 0 0 0 40" fill="none" stroke="currentColor" stroke-width="1"/>
                        </pattern>
                        <rect width="100%" height="100%" fill="url(#grid)"/>
                    </svg>
                </div>

                <div id="uploadPrompt" class="relative z-10">
                    <div class="w-16 h-16 bg-slate-100 rounded-2xl flex items-center justify-center mx-auto mb-4">
                        <i class="fas fa-file-import text-2xl text-slate-400"></i>
                    </div>
                    <p class="text-lg font-semibold text-slate-700 mb-1">Drag & drop your file here</p>
                    <p class="text-sm text-slate-400 mb-4">or click anywhere in this area to browse</p>
                    <div class="flex justify-center gap-2">
                        <span class="px-3 py-1.5 bg-slate-100 rounded-lg text-xs font-semibold text-slate-600 border border-slate-200">
                            <i class="fas fa-file-pdf text-red-500 mr-1.5"></i>PDF
                        </span>
                        <span class="px-3 py-1.5 bg-slate-100 rounded-lg text-xs font-semibold text-slate-600 border border-slate-200">
                            <i class="fas fa-file-word text-blue-600 mr-1.5"></i>Word
                        </span>
                    </div>
                </div>

                <div id="fileSelected" class="hidden relative z-10">
                    <div class="w-16 h-16 bg-green-100 rounded-2xl flex items-center justify-center mx-auto mb-4">
                        <i class="fas fa-check-circle text-2xl text-green-600"></i>
                    </div>
                    <p id="fileName" class="text-lg font-semibold text-slate-800 mb-1"></p>
                    <p id="fileSize" class="text-sm text-slate-500 mb-3"></p>
                    <button id="removeFile" class="px-4 py-2 bg-red-50 hover:bg-red-100 text-red-600 rounded-lg text-sm font-semibold transition-colors border border-red-200">
                        <i class="fas fa-trash-alt mr-1.5"></i>Remove File
                    </button>
                </div>
            </div>

            <!-- Focus Input -->
            <div class="mb-6">
                <label class="flex items-center gap-2 text-sm font-bold text-slate-700 mb-2">
                    <div class="w-6 h-6 bg-purple-100 rounded-lg flex items-center justify-center">
                        <i class="fas fa-crosshairs text-purple-600 text-xs"></i>
                    </div>
                    Specific Focus <span class="text-slate-400 font-normal">(Optional)</span>
                </label>
                <div class="relative">
                    <input 
                        type="text" 
                        id="focusInput" 
                        placeholder="Leave empty for 'Architectural Basics' extraction, or type: 'concrete grade and rebar spacing'"
                        class="w-full px-4 py-3.5 pl-11 bg-slate-50 border border-slate-200 rounded-xl focus:ring-2 focus:ring-blue-500 focus:border-transparent outline-none transition-all text-sm font-medium"
                    >
                    <i class="fas fa-search absolute left-4 top-1/2 -translate-y-1/2 text-slate-400"></i>
                </div>
                <div class="flex items-center gap-2 mt-2">
                    <i class="fas fa-info-circle text-slate-400 text-xs"></i>
                    <p class="text-xs text-slate-500">
                        <strong class="text-slate-700">Default Mode:</strong> Empty input extracts Project Name, Client, Dates, Materials, Budget, Structural Info, Location, Drawing Info
                    </p>
                </div>
            </div>

            <!-- Extract Button -->
            <button id="extractBtn" disabled class="w-full py-4 bg-slate-900 hover:bg-slate-800 text-white font-bold rounded-xl shadow-lg shadow-slate-900/20 transition-all disabled:opacity-40 disabled:cursor-not-allowed disabled:shadow-none flex items-center justify-center gap-2 text-sm tracking-wide">
                <i class="fas fa-magic"></i>
                <span id="btnText">EXTRACT DATA WITH AI</span>
            </button>
        </div>

        <!-- Loading State -->
        <div id="loadingState" class="hidden bg-white rounded-2xl shadow-xl shadow-slate-200/50 p-12 text-center border border-slate-100">
            <div class="spinner mx-auto mb-6"></div>
            <h3 class="text-lg font-bold text-slate-800 mb-2">Analyzing Document...</h3>
            <p id="loadingText" class="text-sm text-slate-500 mb-6">Reading file and sending to AI</p>
            <div class="max-w-sm mx-auto">
                <div class="h-2 bg-slate-100 rounded-full overflow-hidden">
                    <div id="progressBar" class="h-full bg-blue-600 rounded-full transition-all duration-700" style="width: 0%"></div>
                </div>
            </div>
        </div>

        <!-- Results Section -->
        <div id="resultsSection" class="hidden">
            <div class="flex items-center justify-between mb-4">
                <div id="modeBadge" class="px-4 py-2 rounded-xl text-sm font-bold flex items-center gap-2 border"></div>
                <div class="flex gap-2">
                    <button id="copyJson" class="px-4 py-2 bg-slate-100 hover:bg-slate-200 text-slate-700 rounded-xl text-sm font-semibold transition-colors border border-slate-200">
                        <i class="fas fa-copy mr-1.5"></i>Copy JSON
                    </button>
                    <button id="downloadJson" class="px-4 py-2 bg-slate-900 hover:bg-slate-800 text-white rounded-xl text-sm font-semibold transition-colors shadow-lg shadow-slate-900/20">
                        <i class="fas fa-download mr-1.5"></i>Download JSON
                    </button>
                </div>
            </div>
            <div id="resultsContainer" class="space-y-4"></div>
        </div>

        <!-- Error State -->
        <div id="errorState" class="hidden bg-red-50 border border-red-200 rounded-2xl p-6">
            <div class="flex items-start gap-3">
                <div class="w-10 h-10 bg-red-100 rounded-xl flex items-center justify-center flex-shrink-0">
                    <i class="fas fa-exclamation-triangle text-red-600"></i>
                </div>
                <div>
                    <h3 class="font-bold text-red-800 mb-1">Extraction Error</h3>
                    <p id="errorMessage" class="text-sm text-red-600 leading-relaxed"></p>
                </div>
            </div>
        </div>
    </main>

    <!-- Footer -->
    <footer class="max-w-6xl mx-auto px-6 py-8 text-center border-t border-slate-200 mt-8">
        <p class="text-sm text-slate-500 mb-1">Architecture AI Platform v1.0.0</p>
        <p class="text-xs text-slate-400">Local Desktop Application &bull; Documents never leave your machine</p>
    </footer>

    <script>
        // ═══════════════════════════════════════════════════════════════
        // FRONTEND JAVASCRIPT
        // ═══════════════════════════════════════════════════════════════
        const dropZone = document.getElementById('dropZone');
        const fileInput = document.getElementById('fileInput');
        const uploadPrompt = document.getElementById('uploadPrompt');
        const fileSelected = document.getElementById('fileSelected');
        const fileNameEl = document.getElementById('fileName');
        const fileSizeEl = document.getElementById('fileSize');
        const removeFileBtn = document.getElementById('removeFile');
        const focusInput = document.getElementById('focusInput');
        const extractBtn = document.getElementById('extractBtn');
        const btnText = document.getElementById('btnText');
        const loadingState = document.getElementById('loadingState');
        const loadingText = document.getElementById('loadingText');
        const progressBar = document.getElementById('progressBar');
        const resultsSection = document.getElementById('resultsSection');
        const resultsContainer = document.getElementById('resultsContainer');
        const modeBadge = document.getElementById('modeBadge');
        const errorState = document.getElementById('errorState');
        const errorMessage = document.getElementById('errorMessage');
        const apiStatus = document.getElementById('apiStatus');
        const downloadJsonBtn = document.getElementById('downloadJson');
        const copyJsonBtn = document.getElementById('copyJson');

        let currentFile = null;
        let currentResult = null;

        // Check API status on load
        fetch('/api/health').then(r => r.json()).then(d => {
            if (d.api_key_configured) {
                apiStatus.className = 'px-3 py-1.5 rounded-full text-xs font-semibold bg-green-500/15 text-green-400 border border-green-500/20 flex items-center gap-1.5';
                apiStatus.innerHTML = '<span class="w-1.5 h-1.5 rounded-full bg-green-400"></span>API Ready';
            } else {
                apiStatus.className = 'px-3 py-1.5 rounded-full text-xs font-semibold bg-red-500/15 text-red-400 border border-red-500/20 flex items-center gap-1.5';
                apiStatus.innerHTML = '<span class="w-1.5 h-1.5 rounded-full bg-red-400"></span>API Key Missing';
            }
        }).catch(() => {
            apiStatus.className = 'px-3 py-1.5 rounded-full text-xs font-semibold bg-red-500/15 text-red-400 border border-red-500/20 flex items-center gap-1.5';
            apiStatus.innerHTML = '<span class="w-1.5 h-1.5 rounded-full bg-red-400"></span>Server Offline';
        });

        // Drag & Drop handlers
        dropZone.addEventListener('click', () => fileInput.click());
        dropZone.addEventListener('dragover', e => { e.preventDefault(); dropZone.classList.add('dragover'); });
        dropZone.addEventListener('dragleave', () => dropZone.classList.remove('dragover'));
        dropZone.addEventListener('drop', e => {
            e.preventDefault();
            dropZone.classList.remove('dragover');
            if (e.dataTransfer.files.length) handleFile(e.dataTransfer.files[0]);
        });
        fileInput.addEventListener('change', e => { if (e.target.files.length) handleFile(e.target.files[0]); });
        removeFileBtn.addEventListener('click', e => { e.stopPropagation(); clearFile(); });

        function handleFile(file) {
            const exts = ['pdf', 'docx', 'doc'];
            const ext = file.name.split('.').pop().toLowerCase();
            if (!exts.includes(ext)) { showError('Invalid file type. Only PDF and Word documents are supported.'); return; }

            currentFile = file;
            fileNameEl.textContent = file.name;
            fileSizeEl.textContent = formatFileSize(file.size);
            uploadPrompt.classList.add('hidden');
            fileSelected.classList.remove('hidden');
            dropZone.classList.add('has-file');
            extractBtn.disabled = false;
            hideError();
        }

        function clearFile() {
            currentFile = null;
            fileInput.value = '';
            uploadPrompt.classList.remove('hidden');
            fileSelected.classList.add('hidden');
            dropZone.classList.remove('has-file');
            extractBtn.disabled = true;
        }

        function formatFileSize(bytes) {
            if (bytes === 0) return '0 Bytes';
            const k = 1024;
            const sizes = ['Bytes', 'KB', 'MB', 'GB'];
            const i = Math.floor(Math.log(bytes) / Math.log(k));
            return parseFloat((bytes / Math.pow(k, i)).toFixed(2)) + ' ' + sizes[i];
        }

        function showError(msg) {
            errorState.classList.remove('hidden');
            errorMessage.textContent = msg;
            resultsSection.classList.add('hidden');
        }
        function hideError() { errorState.classList.add('hidden'); }

        // Extract button handler
        extractBtn.addEventListener('click', async () => {
            if (!currentFile) return;
            const focus = focusInput.value.trim();

            hideError();
            resultsSection.classList.add('hidden');
            loadingState.classList.remove('hidden');
            extractBtn.disabled = true;
            btnText.textContent = 'PROCESSING...';

            loadingText.textContent = focus 
                ? `Focusing on: "${focus.substring(0, 50)}${focus.length > 50 ? '...' : ''}"` 
                : 'Extracting Architectural Basics...';

            // Progress animation
            let progress = 0;
            const progressInterval = setInterval(() => {
                progress += Math.random() * 12;
                if (progress > 90) progress = 90;
                progressBar.style.width = progress + '%';
            }, 400);

            try {
                const formData = new FormData();
                formData.append('file', currentFile);
                formData.append('focus', focus);

                const res = await fetch('/api/extract', { method: 'POST', body: formData });
                const result = await res.json();

                clearInterval(progressInterval);
                progressBar.style.width = '100%';

                if (res.ok && !result.error) {
                    currentResult = result;
                    setTimeout(() => {
                        loadingState.classList.add('hidden');
                        renderResults(result, focus);
                    }, 300);
                } else {
                    throw new Error(result.error || 'Unknown error occurred');
                }
            } catch (err) {
                clearInterval(progressInterval);
                loadingState.classList.add('hidden');
                showError(err.message);
                extractBtn.disabled = false;
                btnText.textContent = 'EXTRACT DATA WITH AI';
            }
        });

        function renderResults(data, focus) {
            resultsSection.classList.remove('hidden');
            resultsContainer.innerHTML = '';
            extractBtn.disabled = false;
            btnText.textContent = 'EXTRACT DATA WITH AI';

            const isCustom = !!focus;
            modeBadge.className = isCustom 
                ? 'px-4 py-2 rounded-xl text-sm font-bold flex items-center gap-2 border bg-purple-50 text-purple-700 border-purple-200'
                : 'px-4 py-2 rounded-xl text-sm font-bold flex items-center gap-2 border bg-blue-50 text-blue-700 border-blue-200';
            modeBadge.innerHTML = isCustom 
                ? `<i class="fas fa-crosshairs"></i> Custom Focus: "${focus.substring(0, 35)}${focus.length > 35 ? '...' : ''}"`
                : '<i class="fas fa-list-check"></i> Architectural Basics (Default)';

            if (isCustom) renderCustomResults(data);
            else renderDefaultResults(data);

            resultsSection.scrollIntoView({ behavior: 'smooth', block: 'start' });
        }

        function renderDefaultResults(data) {
            const sections = [
                { title: 'Project Overview', icon: 'fa-building', color: 'blue', 
                  items: [['Project Name', data.project_name], ['Client', data.client], ['Contractor', data.contractor], ['Architect', data.architect], ['Engineer', data.engineer]] },
                { title: 'Project Dates', icon: 'fa-calendar-alt', color: 'green',
                  items: data.project_dates ? [['Start Date', data.project_dates.start_date], ['Completion Date', data.project_dates.completion_date], ['Drawing Date', data.project_dates.drawing_date]] : [] },
                { title: 'Budget Information', icon: 'fa-dollar-sign', color: 'emerald',
                  items: data.budget_info ? [['Total Budget', data.budget_info.total_budget], ['Currency', data.budget_info.currency], ['Cost Breakdown', data.budget_info.cost_breakdown]] : [] },
                { title: 'Structural Information', icon: 'fa-drafting-compass', color: 'indigo',
                  items: data.structural_info ? [['Building Type', data.structural_info.building_type], ['Foundation', data.structural_info.foundation_type], ['Framing', data.structural_info.framing_type], ['Roof', data.structural_info.roof_type]] : [] },
                { title: 'Location', icon: 'fa-map-marker-alt', color: 'rose',
                  items: data.location ? [['Address', data.location.address], ['City', data.location.city], ['State', data.location.state], ['ZIP Code', data.location.zip_code]] : [] },
                { title: 'Drawing Information', icon: 'fa-ruler-combined', color: 'cyan',
                  items: data.drawing_info ? [['Sheet Number', data.drawing_info.sheet_number], ['Scale', data.drawing_info.drawing_scale], ['Revision', data.drawing_info.revision]] : [] }
            ];

            // Materials (special tag display)
            if (data.main_materials && data.main_materials.length) {
                const card = createCard('Main Materials', 'fa-cubes', 'amber');
                const div = document.createElement('div');
                div.className = 'flex flex-wrap gap-2';
                data.main_materials.forEach(m => {
                    const tag = document.createElement('span');
                    tag.className = 'px-3 py-1.5 bg-amber-50 text-amber-700 rounded-lg text-xs font-semibold border border-amber-200';
                    tag.textContent = m;
                    div.appendChild(tag);
                });
                card.appendChild(div);
                resultsContainer.appendChild(card);
            }

            sections.forEach(s => {
                if (s.items.length && s.items.some(([_, v]) => v)) {
                    const card = createCard(s.title, s.icon, s.color);
                    card.innerHTML += createGrid(s.items);
                    resultsContainer.appendChild(card);
                }
            });

            // Metadata
            if (data._meta) {
                const meta = createCard('Extraction Metadata', 'fa-cogs', 'slate');
                meta.innerHTML += createGrid([
                    ['Model', data._meta.model_used],
                    ['Mode', data._meta.extraction_mode],
                    ['Tokens Used', data._meta.tokens_used?.toLocaleString()],
                ]);
                resultsContainer.appendChild(meta);
            }

            // Raw JSON
            const raw = createCard('Raw JSON Data', 'fa-code', 'slate');
            const pre = document.createElement('pre');
            pre.className = 'bg-slate-900 text-slate-300 p-4 rounded-xl text-xs overflow-x-auto font-mono leading-relaxed';
            pre.textContent = JSON.stringify(data, null, 2);
            raw.appendChild(pre);
            resultsContainer.appendChild(raw);
        }

        function renderCustomResults(data) {
            const summary = createCard('Extraction Summary', 'fa-clipboard-list', 'purple');
            summary.innerHTML += `
                <div class="space-y-3">
                    <div class="flex items-center gap-2">
                        <span class="text-xs font-semibold text-slate-500 uppercase tracking-wider w-20">Focus:</span>
                        <span class="text-sm text-slate-800 font-medium">${data.extraction_focus || 'N/A'}</span>
                    </div>
                    <div class="flex items-center gap-2">
                        <span class="text-xs font-semibold text-slate-500 uppercase tracking-wider w-20">Confidence:</span>
                        <span class="px-2.5 py-1 rounded-lg text-xs font-bold ${getConfidenceClass(data.confidence)}">${data.confidence || 'N/A'}</span>
                    </div>
                    <p class="text-sm text-slate-600 leading-relaxed">${data.summary || 'No summary available.'}</p>
                </div>
            `;
            resultsContainer.appendChild(summary);

            if (data.found_items && data.found_items.length) {
                const items = createCard(`Found Items (${data.found_items.length})`, 'fa-search', 'blue');
                data.found_items.forEach((item, i) => {
                    const div = document.createElement('div');
                    div.className = 'border-l-4 border-blue-500 pl-4 py-3 mb-3 bg-blue-50/40 rounded-r-xl';
                    div.innerHTML = `
                        <div class="font-bold text-sm text-slate-800 mb-1">${i+1}. ${item.item || 'Unnamed Item'}</div>
                        <div class="text-sm text-slate-700 mb-1"><span class="font-semibold text-slate-500">Value:</span> ${item.value || 'N/A'}</div>
                        <div class="text-xs text-slate-500 mb-1"><span class="font-semibold">Context:</span> ${item.context || 'N/A'}</div>
                        <div class="text-xs text-slate-400"><i class="fas fa-map-pin mr-1"></i>${item.location_hint || 'Location unknown'}</div>
                    `;
                    items.appendChild(div);
                });
                resultsContainer.appendChild(items);
            } else {
                const empty = createCard('Found Items', 'fa-search', 'gray');
                empty.innerHTML += '<p class="text-sm text-slate-500 italic">No items found matching your focus criteria.</p>';
                resultsContainer.appendChild(empty);
            }

            if (data.additional_notes) {
                const notes = createCard('Additional Notes', 'fa-sticky-note', 'amber');
                notes.innerHTML += `<p class="text-sm text-slate-600 leading-relaxed">${data.additional_notes}</p>`;
                resultsContainer.appendChild(notes);
            }

            // Metadata
            if (data._meta) {
                const meta = createCard('Extraction Metadata', 'fa-cogs', 'slate');
                meta.innerHTML += createGrid([
                    ['Model', data._meta.model_used],
                    ['Mode', data._meta.extraction_mode],
                    ['Tokens Used', data._meta.tokens_used?.toLocaleString()],
                ]);
                resultsContainer.appendChild(meta);
            }

            // Raw JSON
            const raw = createCard('Raw JSON Data', 'fa-code', 'slate');
            const pre = document.createElement('pre');
            pre.className = 'bg-slate-900 text-slate-300 p-4 rounded-xl text-xs overflow-x-auto font-mono leading-relaxed';
            pre.textContent = JSON.stringify(data, null, 2);
            raw.appendChild(pre);
            resultsContainer.appendChild(raw);
        }

        function createCard(title, icon, color) {
            const card = document.createElement('div');
            card.className = 'result-card bg-white rounded-2xl shadow-lg shadow-slate-200/40 border border-slate-100 p-6';
            card.innerHTML = `
                <div class="flex items-center gap-3 mb-4 pb-3 border-b border-slate-100">
                    <div class="w-8 h-8 bg-${color}-100 rounded-lg flex items-center justify-center">
                        <i class="fas ${icon} text-${color}-600 text-sm"></i>
                    </div>
                    <h3 class="font-bold text-slate-800 text-sm">${title}</h3>
                </div>
            `;
            return card;
        }

        function createGrid(items) {
            let html = '<div class="grid grid-cols-1 md:grid-cols-2 gap-3">';
            items.forEach(([label, value]) => {
                html += `
                    <div class="bg-slate-50 rounded-xl p-3 border border-slate-100">
                        <div class="text-[10px] font-bold text-slate-400 uppercase tracking-wider mb-1">${label}</div>
                        <div class="text-sm font-semibold text-slate-800">${value || '<span class="text-slate-400 italic font-normal">Not found</span>'}</div>
                    </div>
                `;
            });
            html += '</div>';
            return html;
        }

        function getConfidenceClass(c) {
            if (!c) return 'bg-gray-100 text-gray-600';
            const val = c.toLowerCase();
            if (val === 'high') return 'bg-green-100 text-green-700 border border-green-200';
            if (val === 'medium') return 'bg-yellow-100 text-yellow-700 border border-yellow-200';
            if (val === 'low') return 'bg-red-100 text-red-700 border border-red-200';
            return 'bg-gray-100 text-gray-600';
        }

        // Download JSON
        downloadJsonBtn.addEventListener('click', () => {
            if (!currentResult) return;
            const blob = new Blob([JSON.stringify(currentResult, null, 2)], { type: 'application/json' });
            const url = URL.createObjectURL(blob);
            const a = document.createElement('a');
            a.href = url;
            a.download = `architecture_ai_extract_${new Date().toISOString().slice(0,19).replace(/:/g,'-')}.json`;
            document.body.appendChild(a);
            a.click();
            document.body.removeChild(a);
            URL.revokeObjectURL(url);
        });

        // Copy JSON to clipboard
        copyJsonBtn.addEventListener('click', () => {
            if (!currentResult) return;
            navigator.clipboard.writeText(JSON.stringify(currentResult, null, 2)).then(() => {
                copyJsonBtn.innerHTML = '<i class="fas fa-check mr-1.5"></i>Copied!';
                setTimeout(() => {
                    copyJsonBtn.innerHTML = '<i class="fas fa-copy mr-1.5"></i>Copy JSON';
                }, 2000);
            });
        });
    </script>
</body>
</html>
"""

# =============================================================================
# SECTION 7: FLASK API ROUTES
# =============================================================================
@app.route('/')
def index():
    """Serve the main frontend page."""
    return render_template_string(HTML_TEMPLATE)

@app.route('/api/extract', methods=['POST'])
def api_extract():
    """
    Handle file upload and AI extraction.

    Form Data:
        - file: The uploaded PDF or Word document
        - focus: Optional custom extraction focus (empty = default mode)

    Returns:
        JSON with extracted data or error message
    """
    # Validate file upload
    if 'file' not in request.files:
        return jsonify({"error": "No file provided. Please upload a PDF or Word document."}), 400

    file = request.files['file']
    user_focus = request.form.get('focus', '').strip()

    if file.filename == '':
        return jsonify({"error": "No file selected. Please choose a file."}), 400

    if not allowed_file(file.filename):
        return jsonify({"error": f"Invalid file type. Supported: {', '.join(ALLOWED_EXTENSIONS)}"}), 400

    # Save uploaded file temporarily
    filename = secure_filename(file.filename)
    file_path = UPLOAD_FOLDER / filename
    file.save(file_path)

    try:
        # Extract text from document
        document_text = extract_text(str(file_path), filename)

        if document_text.startswith("ERROR"):
            file_path.unlink(missing_ok=True)
            return jsonify({"error": document_text}), 500

        if not document_text or len(document_text.strip()) < 50:
            file_path.unlink(missing_ok=True)
            return jsonify({"error": "Could not extract meaningful text. The document may be scanned/image-based or password-protected."}), 400

        # Send to OpenAI for extraction
        result = extract_with_openai(document_text, user_focus)

        # Clean up temp file
        file_path.unlink(missing_ok=True)

        return jsonify(result)

    except Exception as e:
        file_path.unlink(missing_ok=True)
        return jsonify({"error": f"Server error: {str(e)}"}), 500

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint - returns API status."""
    return jsonify({
        "status": "ok",
        "app_name": APP_NAME,
        "app_version": APP_VERSION,
        "api_key_configured": bool(OPENAI_API_KEY),
        "supported_formats": list(ALLOWED_EXTENSIONS)
    })

# =============================================================================
# SECTION 8: DESKTOP APP BOOTSTRAP
# =============================================================================
def open_browser_delayed():
    """Open default browser after server starts."""
    time.sleep(2.0)
    url = f"http://localhost:{APP_PORT}"
    print(f"\n🌐 Opening browser at {url}")
    webbrowser.open(url)

def run_desktop_app():
    """Main entry point - runs the complete desktop application."""
    print("\n" + "=" * 65)
    print("  🏗️  ARCHITECTURE AI PLATFORM - Desktop Application")
    print("=" * 65)
    print("\n  📋 Features:")
    print("     • Extract Architectural Basics (Default Mode)")
    print("     • Custom Focus Deep Extraction")
    print("     • PDF & Word document support")
    print("     • Powered by OpenAI GPT-4o-mini")
    print("     • Auto-installs missing dependencies")
    print("\n" + "=" * 65)

    if not OPENAI_API_KEY:
        print("\n  ⚠️  WARNING: OPENAI_API_KEY not set!")
        print("  ┌─────────────────────────────────────────────────────────────┐")
        print("  │  Windows CMD:  set OPENAI_API_KEY=sk-your-key-here          │")
        print("  │  Windows PS:    $env:OPENAI_API_KEY='sk-your-key-here'      │")
        print("  │  Mac/Linux:     export OPENAI_API_KEY=sk-your-key-here      │")
        print("  └─────────────────────────────────────────────────────────────┘")
        print("  Get your key: https://platform.openai.com/api-keys\n")
    else:
        masked = OPENAI_API_KEY[:8] + "..." + OPENAI_API_KEY[-4:]
        print(f"\n  ✅ OpenAI API Key configured: {masked}")

    print(f"\n  🚀 Starting server at http://localhost:{APP_PORT}")
    print("  📂 Press CTRL+C to stop the server\n")

    # Open browser in background thread
    browser_thread = threading.Thread(target=open_browser_delayed, daemon=True)
    browser_thread.start()

    # Run Flask server (production-ready settings)
    app.run(
        host=APP_HOST,
        port=APP_PORT,
        debug=False,
        use_reloader=False,
        threaded=True
    )

if __name__ == '__main__':
    run_desktop_app()
